import os
import sys
import ollama
from pathlib import Path
from PyPDF2 import PdfReader
import markdown
import docx
import subprocess
import tempfile
import json
import glob
import numpy as np
from datetime import datetime
import threading
import queue
from concurrent.futures import ThreadPoolExecutor, as_completed
import time
import socket
import requests
import chromadb
import shutil
from vram_monitor import VRAMMonitor, monitor_distributed_nodes
from gpu_controller import GPUController
from intelligent_gpu_router import IntelligentGPURouter
from adaptive_parallelism import AdaptiveParallelismStrategy

# 🚀 AVAILABLE COMMANDS:
COMMANDS = """
   📖 open_pdf <file>   → Process a single PDF file
   📂 open_dir <dir>    → Process all PDFs in a directory
   💬 chat              → Chat with processed PDFs
   📊 list_docs         → List all processed documents
   🔍 check_deps        → Check for required dependencies
   🌐 discover_nodes    → Auto-discover Ollama nodes on local network
   ➕ add_node <url>    → Manually add an Ollama node (e.g., http://192.168.1.100:11434)
   ➖ remove_node <url> → Remove an Ollama node from the pool
   📋 list_nodes        → List all configured Ollama nodes
   🔬 verify_models     → Check which models are available on each node
   ⚖️  lb_stats          → Show load balancer statistics
   🎯 set_routing <strategy> → Set routing: adaptive, round_robin, least_loaded, lowest_latency
   🖥️  vram_report       → Show detailed VRAM usage report
   🚀 force_gpu <model> → Force model to GPU on all capable nodes
   🎯 gpu_status        → Show intelligent GPU routing status
   🧠 gpu_route <model> → Show routing decision for a model
   🔧 gpu_optimize      → Trigger intelligent GPU optimization
   ✅ gpu_check <model> → Check which nodes can fit a model
   📚 gpu_models        → List all known models and sizes
   🗑️  unload_model <model> → Unload a specific model from memory
   🧹 cleanup_models    → Unload all non-priority models
   🔀 parallelism_report → Show adaptive parallelism analysis
   🧹 clear_cache       → Clear embedding cache (keeps documents)
   🗑️  clear_db          → Clear ChromaDB vector store (removes all documents)
   ❌ exit              → Quit the program
"""

# 🔥 AI MODELS
EMBEDDING_MODEL = "mxbai-embed-large"
CHAT_MODEL = "llama3.1:latest"

# 🚀 MODEL CACHING CONFIGURATION
# Keep models in VRAM for faster inference (prevents reloading)
EMBEDDING_KEEP_ALIVE = "1h"   # Embedding model used frequently for chunking/search
CHAT_KEEP_ALIVE = "15m"       # Chat model used less frequently

# 📊 RAG CONFIGURATION
# Retrieval settings for chat
RETRIEVAL_TOP_K = 10          # Number of chunks to retrieve (default: 10)
RETRIEVAL_MIN_SIMILARITY = 0.3  # Minimum similarity score (0.0-1.0)
CHUNKS_TO_SHOW = 10           # Number of source chunks to display (show all retrieved)

# Acceptable model variations (allows flexible matching)
ACCEPTABLE_EMBEDDING_MODELS = [
    "mxbai-embed-large",
    "mxbai-embed-large:latest",
    "nomic-embed-text",
    "nomic-embed-text:latest",
    "all-minilm",
    "all-minilm:latest",
    "bge-large",
    "bge-large:latest",
]

ACCEPTABLE_CHAT_MODELS = [
    "llama3.1",
    "llama3.1:latest",
    "llama3.1:8b",
    "llama3.1:70b",
    "llama3.2",
    "llama3.2:latest",
    "llama3.2:3b",
    "llama3",
    "llama3:latest",
    "llama3:8b",
    "llama3:70b",
    "mistral",
    "mistral:latest",
    "mixtral",
    "mixtral:latest",
    "qwen",
    "qwen:latest",
    "qwen2.5",
    "phi3",
]

# 🌐 OLLAMA LOAD BALANCER CONFIGURATION
# Add multiple Ollama instances here (host:port)
OLLAMA_INSTANCES = [
    "http://localhost:11434",  # Default local instance
    # Additional nodes will be auto-discovered or loaded from ollama_nodes.json
]

class OllamaLoadBalancer:
    """Load balancer for distributing requests across multiple Ollama instances with intelligent routing."""

    def __init__(self, instances, skip_init_checks=False):
        self.instances = list(instances)  # Make a copy
        self.current_index = 0
        self.lock = threading.Lock()
        self.skip_init_checks = skip_init_checks  # Skip network checks during init
        self.instance_stats = {
            inst: {
                "requests": 0,
                "errors": 0,
                "total_time": 0,
                "latency": None,  # Network latency in ms
                "concurrent_requests": 0,  # Current active requests
                "health_score": 100.0,  # 0-100 health score
                "last_check": None,
                "has_gpu": None,  # GPU availability
                "gpu_memory_gb": 0,  # GPU memory in GB
                "is_local": inst.startswith("http://localhost") or inst.startswith("http://127.0.0.1"),
                "force_cpu": False  # Manual CPU override
            } for inst in instances
        }
        self.nodes_file = KB_DIR / "ollama_nodes.json"
        self.routing_strategy = "adaptive"  # Options: round_robin, least_loaded, lowest_latency, adaptive

        # GPU optimization settings
        self.gpu_controller = GPUController()
        self.auto_optimize_gpu = False  # Disable automatic GPU optimization (use manual gpu_optimize command)
        self.gpu_priority_models = ["mxbai-embed-large", "llama3.1"]  # Models to prefer GPU
        self.optimization_thread = None
        self.optimization_running = False

        # Node availability cache (to avoid checking every request)
        self.availability_cache = {}  # {node_url: (is_available, timestamp)}
        self.availability_cache_ttl = 60  # Cache for 60 seconds

        # Adaptive parallelism strategy
        self.parallelism_strategy = AdaptiveParallelismStrategy(self)
        self.auto_adaptive_mode = True  # Automatically choose parallel vs sequential

        self._load_nodes_from_disk()

        # Auto-discover nodes if only localhost is configured and no saved nodes
        if not skip_init_checks and len(self.instances) == 1 and self.instances[0] in ["http://localhost:11434", "http://127.0.0.1:11434"]:
            if not self.nodes_file.exists() or len(self._load_saved_nodes_list()) == 0:
                print("🔍 No additional nodes found. Run 'discover_nodes' to find Ollama instances on your network.")

        # Skip initial network checks if requested (for MCP server or testing)
        if not skip_init_checks:
            self._measure_initial_latencies()

            # Start background GPU optimization
            if self.auto_optimize_gpu:
                self._start_gpu_optimization()

    def _load_saved_nodes_list(self):
        """Helper to check saved nodes without modifying instances."""
        if self.nodes_file.exists():
            try:
                with open(self.nodes_file, 'r') as f:
                    return json.load(f)
            except:
                pass
        return []

    def _load_nodes_from_disk(self):
        """Load saved nodes from disk."""
        if self.nodes_file.exists():
            try:
                with open(self.nodes_file, 'r') as f:
                    saved_nodes = json.load(f)
                    for node in saved_nodes:
                        # Handle both old format (string) and new format (dict)
                        if isinstance(node, str):
                            node_url = node
                            force_cpu = False
                        else:
                            node_url = node.get("url")
                            force_cpu = node.get("force_cpu", False)

                        if node_url not in self.instances:
                            # Add saved nodes without checking (will be checked at runtime)
                            print(f"📋 Loading saved node: {node_url}{' (forced CPU)' if force_cpu else ''}")
                            self.instances.append(node_url)
                            is_local = node_url.startswith("http://localhost") or node_url.startswith("http://127.0.0.1")
                            self.instance_stats[node_url] = {
                                "requests": 0,
                                "errors": 0,
                                "total_time": 0,
                                "latency": None,
                                "concurrent_requests": 0,
                                "health_score": 50,  # Default score, will be updated
                                "last_check": None,
                                "has_gpu": False if force_cpu else None,  # Force CPU if configured
                                "gpu_memory_gb": 0,
                                "is_gpu_loaded": False,
                                "is_local": is_local,
                                "force_cpu": force_cpu  # Store config
                            }
            except Exception as e:
                print(f"⚠️ Error loading saved nodes: {e}")

    def _save_nodes_to_disk(self):
        """Save current nodes to disk."""
        try:
            with open(self.nodes_file, 'w') as f:
                json.dump(self.instances, f, indent=2)
        except Exception as e:
            print(f"⚠️ Error saving nodes: {e}")

    def _measure_latency(self, node_url, timeout=2):
        """Measure network latency to a node in milliseconds."""
        try:
            start = time.time()
            response = requests.get(f"{node_url}/api/tags", timeout=timeout)
            latency_ms = (time.time() - start) * 1000
            if response.status_code == 200:
                return latency_ms
        except:
            pass
        return None

    def _detect_gpu(self, node_url):
        """
        Detect if node has GPU and VRAM usage using Ollama's /api/ps endpoint.
        Returns: (has_gpu, gpu_info_str, vram_gb, is_gpu_loaded)
        """
        # Check if this node is forced to CPU mode
        if node_url in self.instance_stats and self.instance_stats[node_url].get("force_cpu", False):
            return False, "CPU (forced)", 0, False

        try:
            # Method 1: Check /api/ps for actual VRAM usage
            try:
                response = requests.get(f"{node_url}/api/ps", timeout=5)
                if response.status_code == 200:
                    ps_data = response.json()
                    models = ps_data.get("models", [])

                    # Check if embedding model is loaded
                    total_vram_bytes = 0
                    embedding_vram_bytes = 0
                    embedding_model_found = False
                    for model_info in models:
                        model_name = model_info.get("name", "")
                        size_vram = model_info.get("size_vram", 0)
                        total_vram_bytes += size_vram

                        # Check if this is our embedding model
                        if any(variant in model_name.lower() for variant in ["mxbai-embed", "nomic-embed", "bge-large", "all-minilm"]):
                            embedding_vram_bytes = size_vram
                            embedding_model_found = True

                    # If embedding model found with VRAM usage, it's using GPU
                    if embedding_vram_bytes > 0:
                        vram_gb = total_vram_bytes / (1024**3)  # Convert to GB
                        return True, f"GPU (model in VRAM: {vram_gb:.1f}GB)", vram_gb, True
                    # If embedding model found but size_vram is 0, it's CPU-only
                    elif embedding_model_found and embedding_vram_bytes == 0:
                        return False, "CPU (model in RAM, not VRAM)", 0, False
                    elif total_vram_bytes > 0:
                        # Other models using GPU, but embedding model not loaded
                        vram_gb = total_vram_bytes / (1024**3)
                        return True, f"GPU available ({vram_gb:.1f}GB used by other models)", vram_gb, False
                    # No models in VRAM, fall through to performance test
            except:
                pass  # Fall back to performance testing

            # Method 2: Performance-based detection (fallback)
            client = ollama.Client(host=node_url)

            # Small embedding test
            start = time.time()
            try:
                test_result = client.embed(model=EMBEDDING_MODEL, input="test")
                small_duration = time.time() - start
            except:
                return None, "Unknown", 0, False

            # Larger batch to stress test
            start = time.time()
            try:
                batch_input = ["This is a longer test input for embedding generation"] * 10
                test_result = client.embed(model=EMBEDDING_MODEL, input=batch_input)
                batch_duration = time.time() - start
                batch_per_item = batch_duration / len(batch_input)
            except:
                batch_per_item = small_duration * 10

            # Performance heuristics
            is_fast_single = small_duration < 0.5
            is_fast_batch = batch_per_item < 0.1

            if is_fast_single and is_fast_batch:
                # Full GPU performance
                vram_estimate = 8 if batch_per_item < 0.05 else 4
                return True, f"GPU (inferred, ~{vram_estimate}GB VRAM)", vram_estimate, True
            elif is_fast_single and not is_fast_batch:
                # GPU with VRAM constraints
                return True, f"GPU (VRAM limited, CPU fallback)", 2, False
            else:
                # CPU-only
                return False, f"CPU only ({small_duration:.2f}s)", 0, False

        except Exception as e:
            return None, f"Detection failed: {str(e)[:40]}", 0, False

    def _measure_initial_latencies(self):
        """Measure initial latency and detect GPU on all nodes."""
        print("📊 Measuring node capabilities...")
        for node in self.instances:
            latency = self._measure_latency(node)
            if latency:
                self.instance_stats[node]["latency"] = latency

                # Detect GPU and VRAM
                has_gpu, gpu_info, vram_gb, is_gpu_loaded = self._detect_gpu(node)
                self.instance_stats[node]["has_gpu"] = has_gpu
                self.instance_stats[node]["gpu_memory_gb"] = vram_gb
                self.instance_stats[node]["is_gpu_loaded"] = is_gpu_loaded

                # Emoji indicators
                if has_gpu and is_gpu_loaded:
                    gpu_emoji = "🚀"  # Full GPU power
                elif has_gpu and not is_gpu_loaded:
                    gpu_emoji = "⚠️"   # GPU with VRAM issues
                elif has_gpu == False:
                    gpu_emoji = "🐢"  # CPU only
                else:
                    gpu_emoji = "❓"  # Unknown

                print(f"   {node}: {latency:.0f}ms {gpu_emoji} {gpu_info}")
            else:
                self.instance_stats[node]["latency"] = 9999
                self.instance_stats[node]["has_gpu"] = False
                self.instance_stats[node]["is_gpu_loaded"] = False
                print(f"   {node}: unreachable")

    def _update_health_score(self, node):
        """Calculate health score with GPU priority, VRAM awareness, latency, error rate, and response time."""
        stats = self.instance_stats[node]

        # Base score
        score = 100.0

        # GPU and VRAM awareness - HEAVILY prioritize GPU nodes
        has_gpu = stats.get("has_gpu")
        is_gpu_loaded = stats.get("is_gpu_loaded", False)
        vram_gb = stats.get("gpu_memory_gb", 0)

        if has_gpu == True:
            if is_gpu_loaded:
                # Full GPU with sufficient VRAM - MASSIVE bonus
                score += 300 + (vram_gb * 20)  # Increased: GPU should always win
            else:
                # GPU exists but NOT being used (size_vram = 0) - treat as CPU
                # This happens when GPU node is configured for CPU-only generation
                score -= 50  # Same penalty as CPU-only nodes
        elif has_gpu == False:
            score -= 100  # CPU nodes heavily penalized (increased from -50)

        # Penalize for high latency (>100ms is bad)
        if stats["latency"]:
            if stats["latency"] > 100:
                score -= min(50, (stats["latency"] - 100) / 10)

        # Penalize for errors
        if stats["requests"] > 0:
            error_rate = stats["errors"] / stats["requests"]
            score -= error_rate * 40

        # Penalize for slow response times
        if stats["requests"] > 0:
            avg_time = stats["total_time"] / stats["requests"]
            if avg_time > 2.0:  # Slower than 2 seconds
                score -= min(30, (avg_time - 2.0) * 10)

        # Penalize for high concurrent load
        load_penalty = stats["concurrent_requests"] * 5
        if has_gpu and is_gpu_loaded:
            # Full GPU can handle more concurrent load
            load_penalty = load_penalty * 0.4
        elif has_gpu and not is_gpu_loaded:
            # VRAM-limited GPU - heavily penalize concurrent load (will fall back to CPU!)
            load_penalty = load_penalty * 1.5
        score -= load_penalty

        stats["health_score"] = max(0, score)
        return stats["health_score"]

    def set_routing_strategy(self, strategy):
        """Set the routing strategy. Options: round_robin, least_loaded, lowest_latency, adaptive"""
        valid_strategies = ["round_robin", "least_loaded", "lowest_latency", "adaptive"]
        if strategy in valid_strategies:
            self.routing_strategy = strategy
            print(f"✅ Routing strategy set to: {strategy}")
        else:
            print(f"❌ Invalid strategy. Choose from: {valid_strategies}")

    def _model_matches(self, available_model, acceptable_models):
        """Check if an available model matches any acceptable model variant."""
        # Normalize the available model name
        available_normalized = available_model.lower().strip()

        for acceptable in acceptable_models:
            acceptable_normalized = acceptable.lower().strip()

            # Direct match
            if available_normalized == acceptable_normalized:
                return True

            # Match without :latest suffix
            if available_normalized.replace(':latest', '') == acceptable_normalized.replace(':latest', ''):
                return True

            # Partial match (e.g., "llama3.1:8b" matches "llama3.1")
            if available_normalized.startswith(acceptable_normalized) or acceptable_normalized in available_normalized:
                return True

        return False

    def _check_model_available(self, node_url, model_name, acceptable_variants=None):
        """Check if a specific model (or acceptable variant) is available on a node."""
        try:
            client = ollama.Client(host=node_url)
            result = client.list()
            if hasattr(result, 'models'):
                models = result.models
                model_names = [model.model if hasattr(model, 'model') else str(model) for model in models]
            else:
                models = result.get('models', [])
                model_names = [model.get('name', 'unknown') for model in models]

            # If acceptable variants provided, check against those
            if acceptable_variants:
                for available_model in model_names:
                    if self._model_matches(available_model, acceptable_variants):
                        return True, available_model
                return False, None
            else:
                # Original behavior - exact match
                for name in model_names:
                    if model_name in name:
                        return True, name
                return False, None
        except:
            return False, None

    def add_node(self, node_url, save=True, check_models=True, optional=False):
        """Add a new Ollama node to the pool.

        Args:
            node_url: URL of the Ollama node
            save: Save to disk
            check_models: Verify required models exist
            optional: If True, add node even if currently offline (will be checked at runtime)
        """
        # Check if this is a duplicate of localhost
        is_localhost_variant = node_url in ["http://localhost:11434", "http://127.0.0.1:11434"]

        # Get local machine IPs to detect duplicates
        import socket
        local_ips = set()
        try:
            hostname = socket.gethostname()
            local_ips.add(socket.gethostbyname(hostname))
        except:
            pass

        # Extract IP from node_url
        node_ip = node_url.replace("http://", "").replace("https://", "").split(":")[0]

        # Check if this node is localhost in disguise
        if node_ip in local_ips:
            # Check if we already have localhost
            has_localhost = any(inst in ["http://localhost:11434", "http://127.0.0.1:11434"]
                               for inst in self.instances)
            if has_localhost and not is_localhost_variant:
                print(f"⚠️  Skipping {node_url} - already have localhost (same machine)")
                return False

        with self.lock:
            if node_url not in self.instances:
                # Test if node is reachable
                try:
                    client = ollama.Client(host=node_url)
                    client.list()  # Test connection

                    # Check if required models are available (with flexible matching)
                    if check_models:
                        has_embedding, embedding_model = self._check_model_available(
                            node_url, EMBEDDING_MODEL, ACCEPTABLE_EMBEDDING_MODELS
                        )
                        has_chat, chat_model = self._check_model_available(
                            node_url, CHAT_MODEL, ACCEPTABLE_CHAT_MODELS
                        )

                        if not has_embedding:
                            print(f"⚠️ Node {node_url} missing compatible embedding model")
                            print(f"   Acceptable: {', '.join(ACCEPTABLE_EMBEDDING_MODELS[:3])}...")
                            print(f"   Run: ssh <host> 'ollama pull {EMBEDDING_MODEL}'")
                            return False
                        else:
                            print(f"   ✅ Found embedding model: {embedding_model}")

                        if not has_chat:
                            print(f"⚠️ Node {node_url} missing compatible chat model (optional)")
                        else:
                            print(f"   ✅ Found chat model: {chat_model}")

                    self.instances.append(node_url)
                    is_local = node_url.startswith("http://localhost") or node_url.startswith("http://127.0.0.1")
                    self.instance_stats[node_url] = {
                        "requests": 0,
                        "errors": 0,
                        "total_time": 0,
                        "latency": None,
                        "concurrent_requests": 0,
                        "health_score": 100.0,
                        "last_check": None,
                        "has_gpu": None,
                        "gpu_memory_gb": 0,
                        "is_gpu_loaded": False,
                        "is_local": is_local
                    }
                    # Measure latency and detect GPU
                    latency = self._measure_latency(node_url)
                    if latency:
                        self.instance_stats[node_url]["latency"] = latency

                    has_gpu, gpu_info, vram_gb, is_gpu_loaded = self._detect_gpu(node_url)
                    self.instance_stats[node_url]["has_gpu"] = has_gpu
                    self.instance_stats[node_url]["gpu_memory_gb"] = vram_gb
                    self.instance_stats[node_url]["is_gpu_loaded"] = is_gpu_loaded

                    if has_gpu and is_gpu_loaded:
                        print(f"   🚀 GPU detected: {gpu_info}")
                    elif has_gpu and not is_gpu_loaded:
                        print(f"   ⚠️  GPU detected (VRAM limited): {gpu_info}")
                    elif has_gpu == False:
                        print(f"   🐢 CPU detected: {gpu_info}")
                    if save:
                        self._save_nodes_to_disk()
                    print(f"✅ Added node: {node_url}")
                    return True
                except Exception as e:
                    if optional:
                        # Add as optional node (offline now, but will check at runtime)
                        print(f"⚠️  Node {node_url} currently offline, adding as optional")
                        self.instances.append(node_url)
                        is_local = node_url.startswith("http://localhost") or node_url.startswith("http://127.0.0.1")
                        self.instance_stats[node_url] = {
                            "requests": 0,
                            "errors": 0,
                            "total_time": 0,
                            "latency": None,
                            "concurrent_requests": 0,
                            "health_score": 0,  # Start with 0 since offline
                            "last_check": None,
                            "has_gpu": None,
                            "gpu_memory_gb": 0,
                            "is_gpu_loaded": False,
                            "is_local": is_local
                        }
                        if save:
                            self._save_nodes_to_disk()
                        return True
                    else:
                        print(f"❌ Failed to add node {node_url}: {e}")
                        return False
            else:
                print(f"ℹ️ Node {node_url} already exists")
                return False

    def remove_node(self, node_url):
        """Remove a node from the pool."""
        with self.lock:
            if node_url in self.instances:
                if len(self.instances) == 1:
                    print("❌ Cannot remove the last node!")
                    return False
                self.instances.remove(node_url)
                if node_url in self.instance_stats:
                    del self.instance_stats[node_url]
                self._save_nodes_to_disk()
                print(f"✅ Removed node: {node_url}")
                return True
            else:
                print(f"⚠️ Node {node_url} not found")
                return False

    def list_nodes(self):
        """List all configured nodes with online/offline status."""
        with self.lock:
            print("\n🌐 Configured Ollama Nodes:")
            print("-" * 60)
            for i, node in enumerate(self.instances, 1):
                stats = self.instance_stats.get(node, {})

                # Check if node is online (don't use cache for list_nodes - show real-time)
                is_online = self._is_node_available(node, timeout=1, use_cache=False)
                if is_online:
                    online_status = "🟢 ONLINE"
                else:
                    online_status = "🔴 OFFLINE"

                # Usage status
                usage_status = "Active" if stats.get("requests", 0) > 0 else "Unused"

                print(f"{i}. {node} - {online_status} - {usage_status}")
                if stats.get("requests", 0) > 0:
                    error_rate = (stats.get("errors", 0) / stats["requests"]) * 100
                    avg_time = stats.get("total_time", 0) / stats["requests"]
                    print(f"   Requests: {stats['requests']}, Errors: {stats.get('errors', 0)} ({error_rate:.1f}%)")
                    print(f"   Avg Response Time: {avg_time:.2f}s")
            print("-" * 60)

    def verify_models_on_nodes(self):
        """Verify which models are available on each node with flexible matching."""
        print("\n🔬 Model Verification:")
        print("-" * 70)

        for node in self.instances:
            print(f"\n🖥️  {node}")
            try:
                client = ollama.Client(host=node)
                result = client.list()

                if hasattr(result, 'models'):
                    models = result.models
                    model_list = [model.model if hasattr(model, 'model') else str(model) for model in models]
                else:
                    models = result.get('models', [])
                    model_list = [model.get('name', 'unknown') for model in models]

                # Check for compatible models using flexible matching
                has_embedding, embedding_found = self._check_model_available(
                    node, EMBEDDING_MODEL, ACCEPTABLE_EMBEDDING_MODELS
                )
                has_chat, chat_found = self._check_model_available(
                    node, CHAT_MODEL, ACCEPTABLE_CHAT_MODELS
                )

                if has_embedding:
                    print(f"   Embedding: ✅ {embedding_found}")
                else:
                    print(f"   Embedding: ❌ None found")
                    print(f"              Acceptable: {', '.join(ACCEPTABLE_EMBEDDING_MODELS[:3])}...")

                if has_chat:
                    print(f"   Chat: ✅ {chat_found}")
                else:
                    print(f"   Chat: ❌ None found")
                    print(f"         Acceptable: {', '.join(ACCEPTABLE_CHAT_MODELS[:3])}...")

                # Show all available models
                print(f"   Total models: {len(model_list)}")
                if model_list:
                    print(f"   All models: {', '.join(model_list[:5])}")
                    if len(model_list) > 5:
                        print(f"               ... and {len(model_list) - 5} more")

            except Exception as e:
                print(f"   ❌ Error: {str(e)}")

        print("-" * 70)

    def discover_nodes(self, require_embedding_model=True):
        """Auto-discover Ollama nodes on the local network with model verification."""
        print("🔍 Scanning local network for Ollama nodes...")
        print("This may take 30-60 seconds...")

        # Get local network range
        hostname = socket.gethostname()
        local_ip = socket.gethostbyname(hostname)
        print(f"📍 Local IP: {local_ip}")

        # Extract network prefix (assumes /24 subnet)
        ip_parts = local_ip.split('.')
        network_prefix = '.'.join(ip_parts[:3])

        discovered = []

        def check_host(ip):
            """Check if Ollama is running on this host."""
            url = f"http://{ip}:11434"
            try:
                response = requests.get(f"{url}/api/tags", timeout=2)
                if response.status_code == 200:
                    return url
            except:
                pass
            return None

        # Scan network in parallel
        with ThreadPoolExecutor(max_workers=50) as executor:
            futures = []
            for i in range(1, 255):
                ip = f"{network_prefix}.{i}"
                futures.append(executor.submit(check_host, ip))

            for future in as_completed(futures):
                result = future.result()
                if result:
                    discovered.append(result)
                    print(f"✅ Found Ollama node: {result}")

        if discovered:
            print(f"\n🎉 Discovered {len(discovered)} Ollama nodes!")
            print(f"Verifying models on discovered nodes...")
            for node in discovered:
                # Add with model checking (will auto-verify compatible models)
                self.add_node(node, check_models=require_embedding_model)
        else:
            print("⚠️ No Ollama nodes found on the network")

        return discovered

    def get_next_instance(self):
        """Round-robin instance selection."""
        with self.lock:
            instance = self.instances[self.current_index]
            self.current_index = (self.current_index + 1) % len(self.instances)
            return instance

    def _is_node_available(self, node_url, timeout=2, use_cache=True):
        """Quick check if node is online and responding."""
        # Check cache first
        if use_cache and node_url in self.availability_cache:
            is_available, timestamp = self.availability_cache[node_url]
            # Use cached result if less than TTL seconds old
            if time.time() - timestamp < self.availability_cache_ttl:
                return is_available

        # Actually check node
        try:
            response = requests.get(f"{node_url}/api/tags", timeout=timeout)
            is_available = response.status_code == 200
        except:
            is_available = False

        # Update cache
        if use_cache:
            self.availability_cache[node_url] = (is_available, time.time())

        return is_available

    def get_available_instances(self):
        """Get list of currently available (online) instances with caching."""
        available = []
        for instance in self.instances:
            if self._is_node_available(instance, use_cache=True):
                available.append(instance)
            else:
                # Mark as offline in stats
                with self.lock:
                    self.instance_stats[instance]["health_score"] = 0
        return available if available else self.instances  # Fallback to all if none available

    def get_best_instance(self):
        """Get instance based on current routing strategy."""
        # Lazy initialization: if we skipped init checks, do them now on first use
        if self.skip_init_checks and not hasattr(self, '_lazy_init_done'):
            self._lazy_init_done = True
            print("🔄 Initializing node stats (first use)...")
            self._measure_initial_latencies()

        # Check available nodes OUTSIDE the lock (involves network calls)
        available_instances = self.get_available_instances()

        with self.lock:

            if not available_instances:
                print("⚠️  Warning: No nodes available, using fallback")
                return self.instances[0] if self.instances else None

            if self.routing_strategy == "round_robin":
                # Round-robin through available nodes only
                for _ in range(len(self.instances)):
                    instance = self.instances[self.current_index]
                    self.current_index = (self.current_index + 1) % len(self.instances)
                    if instance in available_instances:
                        return instance
                return available_instances[0]

            elif self.routing_strategy == "least_loaded":
                # Choose available node with fewest concurrent requests
                return min(available_instances,
                          key=lambda n: self.instance_stats[n]["concurrent_requests"])

            elif self.routing_strategy == "lowest_latency":
                # Choose available node with lowest latency
                viable_nodes = [n for n in available_instances
                                 if self.instance_stats[n]["latency"] is not None
                                 and self.instance_stats[n]["latency"] < 5000]
                if not viable_nodes:
                    return available_instances[0]
                return min(viable_nodes,
                          key=lambda n: self.instance_stats[n]["latency"])

            else:  # adaptive (default)
                # Use health score for adaptive routing (available nodes only)
                best_instance = None
                best_score = -1

                for inst in available_instances:
                    score = self._update_health_score(inst)

                    if score > best_score:
                        best_score = score
                        best_instance = inst

                # Debug: print routing decision
                if best_instance:
                    stats = self.instance_stats[best_instance]
                    gpu_status = "GPU" if stats.get("has_gpu") else "CPU"
                    # Only print for first few requests to avoid spam
                    if stats.get("requests", 0) < 3:
                        print(f"   🎯 Routing to: {best_instance} ({gpu_status}, score={best_score:.0f})")

                return best_instance if best_instance else available_instances[0]

    def record_request(self, instance, duration, error=False):
        """Record request statistics."""
        with self.lock:
            stats = self.instance_stats[instance]
            stats["requests"] += 1
            stats["total_time"] += duration
            if error:
                stats["errors"] += 1

    def embed_distributed(self, model, input_text, keep_alive=None):
        """Generate embedding with automatic failover and concurrent tracking."""
        last_error = None

        # Try all instances
        for _ in range(len(self.instances)):
            instance = self.get_best_instance()

            # Track concurrent request
            with self.lock:
                self.instance_stats[instance]["concurrent_requests"] += 1

            start_time = time.time()

            try:
                # Create client for specific instance
                client = ollama.Client(host=instance)
                # Add keep_alive parameter if provided
                embed_kwargs = {"model": model, "input": input_text}
                if keep_alive:
                    embed_kwargs["keep_alive"] = keep_alive
                result = client.embed(**embed_kwargs)

                duration = time.time() - start_time
                self.record_request(instance, duration, error=False)

                # Update latency if request was fast (good indicator)
                if duration < 3.0:
                    with self.lock:
                        # Running average of latency
                        current_latency = self.instance_stats[instance]["latency"] or duration * 1000
                        self.instance_stats[instance]["latency"] = (current_latency * 0.8 + duration * 1000 * 0.2)

                # Detect VRAM exhaustion: GPU node suddenly running slow
                with self.lock:
                    stats = self.instance_stats[instance]
                    if stats.get("has_gpu") == True and stats.get("is_gpu_loaded") == True:
                        # If a GPU node that should be fast is suddenly slow, it might be VRAM exhaustion
                        if duration > 2.0:  # GPU embedding should be <0.5s
                            print(f"⚠️  {instance} running slow ({duration:.2f}s) - possible VRAM exhaustion")
                            stats["is_gpu_loaded"] = False  # Mark as VRAM-limited
                            stats["health_score"] -= 100  # Heavy penalty

                return result
            except Exception as e:
                duration = time.time() - start_time
                self.record_request(instance, duration, error=True)
                last_error = e
                continue
            finally:
                # Decrease concurrent request counter
                with self.lock:
                    self.instance_stats[instance]["concurrent_requests"] = max(
                        0, self.instance_stats[instance]["concurrent_requests"] - 1
                    )

        # All instances failed
        raise Exception(f"All Ollama instances failed. Last error: {last_error}")

    def chat_distributed(self, model, messages, keep_alive=None):
        """Generate chat response with automatic failover and concurrent tracking."""
        last_error = None

        # Try all instances
        for _ in range(len(self.instances)):
            instance = self.get_best_instance()

            # Track concurrent request
            with self.lock:
                self.instance_stats[instance]["concurrent_requests"] += 1

            start_time = time.time()

            try:
                # Create client for specific instance
                client = ollama.Client(host=instance)
                # Add keep_alive parameter if provided
                chat_kwargs = {"model": model, "messages": messages}
                if keep_alive:
                    chat_kwargs["keep_alive"] = keep_alive
                result = client.chat(**chat_kwargs)

                duration = time.time() - start_time
                self.record_request(instance, duration, error=False)

                # Update latency if request was reasonable
                if duration < 10.0:
                    with self.lock:
                        current_latency = self.instance_stats[instance]["latency"] or duration * 1000
                        self.instance_stats[instance]["latency"] = (current_latency * 0.8 + duration * 1000 * 0.2)

                return result
            except Exception as e:
                duration = time.time() - start_time
                self.record_request(instance, duration, error=True)
                last_error = e
                continue
            finally:
                # Decrease concurrent request counter
                with self.lock:
                    self.instance_stats[instance]["concurrent_requests"] = max(
                        0, self.instance_stats[instance]["concurrent_requests"] - 1
                    )

        # All instances failed
        raise Exception(f"All Ollama instances failed. Last error: {last_error}")

    def embed_batch(self, model, texts, max_workers=None, force_mode=None):
        """
        Embed multiple texts with adaptive parallel/sequential routing.

        Args:
            model: Model name
            texts: List of texts to embed
            max_workers: Manual worker count (overrides adaptive)
            force_mode: 'parallel', 'sequential', or None for adaptive
        """
        batch_size = len(texts)
        results = [None] * batch_size

        # Adaptive mode decision
        if self.auto_adaptive_mode and force_mode is None:
            should_parallel, reasoning = self.parallelism_strategy.should_parallelize(batch_size)

            # Print decision
            print(f"   🔀 Adaptive mode: {reasoning['reason']}")
            print(f"      {reasoning['detail']}")

            if should_parallel:
                mode = "parallel"
            else:
                mode = "sequential"
        else:
            mode = force_mode or "parallel"

        # SEQUENTIAL MODE: Use fastest node only
        if mode == "sequential":
            fastest_node = None
            best_score = -1

            # Find fastest node
            for node in self.instances:
                score = self._update_health_score(node)
                if score > best_score:
                    best_score = score
                    fastest_node = node

            print(f"   ➡️  Sequential mode: Using {fastest_node}")

            # Process sequentially on fastest node
            # Create client once to avoid connection overhead
            client = ollama.Client(host=fastest_node)
            completed = 0

            for i, text in enumerate(texts):
                try:
                    # Use pre-created client for efficiency
                    result = client.embed(model=model, input=text)
                    results[i] = result
                    completed += 1

                    # Progress
                    if completed % 50 == 0 or completed == batch_size:
                        print(f"   Progress: {completed}/{batch_size} embeddings ({completed*100//batch_size}%)")
                except Exception as e:
                    print(f"⚠️ Error embedding text {i}: {e}")

            return results

        # PARALLEL MODE: Distribute across nodes
        else:
            if max_workers is None:
                max_workers = self.parallelism_strategy.get_optimal_workers(batch_size)

            print(f"   🔀 Parallel mode: Using {max_workers} workers across {len(self.instances)} nodes")

            completed = 0

            def embed_single(index, text):
                try:
                    result = self.embed_distributed(model, text)
                    return index, result, None
                except Exception as e:
                    return index, None, e

            with ThreadPoolExecutor(max_workers=max_workers) as executor:
                futures = {executor.submit(embed_single, i, text): i for i, text in enumerate(texts)}

                for future in as_completed(futures):
                    index, result, error = future.result()
                    completed += 1

                    # Show progress every 50 embeddings or on completion
                    if completed % 50 == 0 or completed == batch_size:
                        print(f"   Progress: {completed}/{batch_size} embeddings ({completed*100//batch_size}%)")

                    if error:
                        print(f"⚠️ Error embedding text {index}: {error}")
                    else:
                        results[index] = result

            return results

    def print_stats(self):
        """Print load balancer statistics with routing info."""
        print("\n📊 Ollama Load Balancer Statistics:")
        print(f"Routing Strategy: {self.routing_strategy.upper()}")
        print("-" * 70)

        for inst in self.instances:
            stats = self.instance_stats[inst]

            # Calculate metrics
            if stats["requests"] > 0:
                error_rate = (stats["errors"] / stats["requests"]) * 100
                avg_time = stats["total_time"] / stats["requests"]
            else:
                error_rate = 0
                avg_time = 0

            health_score = self._update_health_score(inst)
            latency = stats["latency"] or 9999

            # Determine status emoji
            if health_score > 80:
                status = "🟢"
            elif health_score > 50:
                status = "🟡"
            else:
                status = "🔴"

            # GPU and VRAM indicators
            has_gpu = stats.get("has_gpu")
            is_gpu_loaded = stats.get("is_gpu_loaded", False)
            vram_gb = stats.get("gpu_memory_gb", 0)

            if has_gpu and is_gpu_loaded:
                gpu_indicator = f" 🚀 GPU (~{vram_gb}GB VRAM)"
            elif has_gpu and not is_gpu_loaded:
                gpu_indicator = " ⚠️  GPU (VRAM limited)"
            elif has_gpu == False:
                gpu_indicator = " 🐢 CPU"
            else:
                gpu_indicator = ""

            print(f"\n{status} {inst}{gpu_indicator}")
            print(f"   Health Score: {health_score:.1f}/100")
            print(f"   Latency: {latency:.0f}ms")
            print(f"   Requests: {stats['requests']}, Errors: {stats['errors']} ({error_rate:.1f}%)")
            if stats["requests"] > 0:
                print(f"   Avg Response: {avg_time:.2f}s")
            print(f"   Concurrent: {stats['concurrent_requests']}")

        print("-" * 70)

    def _start_gpu_optimization(self):
        """Start background thread for automatic GPU optimization."""
        if self.optimization_running:
            return

        self.optimization_running = True
        self.optimization_thread = threading.Thread(
            target=self._gpu_optimization_loop,
            daemon=True,
            name="GPUOptimizer"
        )
        self.optimization_thread.start()
        print("🚀 GPU auto-optimization enabled (background thread)")

    def _gpu_optimization_loop(self):
        """Background loop that periodically optimizes GPU assignments."""
        check_interval = 300  # Check every 5 minutes

        while self.optimization_running:
            try:
                time.sleep(check_interval)

                if not self.auto_optimize_gpu:
                    continue

                print("\n🔧 [GPU Optimizer] Running periodic optimization...")

                # Check each node
                for node_url in self.instances:
                    status = self.gpu_controller.get_model_status(node_url)

                    if 'error' in status:
                        continue

                    # Check if priority models are on CPU when GPU is available
                    stats = self.instance_stats.get(node_url, {})
                    has_gpu = stats.get("has_gpu", False)

                    if not has_gpu:
                        continue  # Skip CPU-only nodes

                    # Check each loaded model
                    for model_info in status.get('models', []):
                        model_name = model_info['name']
                        location = model_info['location']

                        # Check if this is a priority model that should be on GPU
                        is_priority = any(
                            priority in model_name.lower()
                            for priority in self.gpu_priority_models
                        )

                        if is_priority and 'CPU' in location:
                            print(f"⚠️  [GPU Optimizer] {model_name} on CPU at {node_url}, moving to GPU...")
                            result = self.gpu_controller.force_gpu_load(node_url, model_name)
                            print(f"   {result['message']}")

            except Exception as e:
                print(f"⚠️  [GPU Optimizer] Error: {e}")

    def stop_gpu_optimization(self):
        """Stop the background GPU optimization thread."""
        self.optimization_running = False
        if self.optimization_thread:
            self.optimization_thread.join(timeout=5)
        print("🛑 GPU auto-optimization stopped")

    def force_gpu_all_nodes(self, model_name: str):
        """Force a specific model to GPU on all capable nodes."""
        print(f"\n🚀 Forcing {model_name} to GPU on all nodes...")
        results = []

        for node_url in self.instances:
            stats = self.instance_stats.get(node_url, {})
            has_gpu = stats.get("has_gpu", False)

            if not has_gpu:
                print(f"   ⏭️  Skipping {node_url} (no GPU)")
                continue

            print(f"   🔄 Processing {node_url}...")
            result = self.gpu_controller.force_gpu_load(node_url, model_name)
            results.append({
                'node': node_url,
                'result': result
            })
            print(f"      {result['message']}")

        return results

# 📂 File Storage
# Use absolute paths based on script location for MCP server compatibility
_SCRIPT_DIR = Path(__file__).parent.resolve()
PROCESSED_DIR = _SCRIPT_DIR / "converted_files"
PROCESSED_DIR.mkdir(exist_ok=True)

# 📚 Knowledge Base (legacy JSON storage - kept for backwards compatibility)
KB_DIR = _SCRIPT_DIR / "knowledge_base"
KB_DIR.mkdir(exist_ok=True)

# 🗄️ ChromaDB Vector Store (production storage)
CHROMA_DB_DIR = _SCRIPT_DIR / "chroma_db_cli"
CHROMA_DB_DIR.mkdir(exist_ok=True)

# Initialize ChromaDB client and collection
chroma_client = chromadb.PersistentClient(path=str(CHROMA_DB_DIR))
chroma_collection = chroma_client.get_or_create_collection(
    name="documents",
    metadata={"hnsw:space": "cosine"}  # Use cosine similarity for better semantic search
)

# Initialize global load balancer (after KB_DIR is defined)
# Skip init checks when imported as module (for MCP server)
_is_module = __name__ != "__main__"
load_balancer = OllamaLoadBalancer(OLLAMA_INSTANCES, skip_init_checks=_is_module)

# 💾 Index file for tracking processed documents
INDEX_FILE = KB_DIR / "document_index.json"

# 🔄 Cache for embeddings to avoid regenerating
EMBEDDING_CACHE_FILE = KB_DIR / "embedding_cache.json"

def load_embedding_cache():
    """Load the embedding cache from disk."""
    if not EMBEDDING_CACHE_FILE.exists():
        return {}
    try:
        with open(EMBEDDING_CACHE_FILE, 'r') as f:
            return json.load(f)
    except (json.JSONDecodeError, FileNotFoundError):
        return {}

def save_embedding_cache(cache):
    """Save the embedding cache to disk."""
    with open(EMBEDDING_CACHE_FILE, 'w') as f:
        json.dump(cache, f)

def get_cached_embedding(text, use_load_balancer=True):
    """Get embedding from cache or generate new one."""
    import hashlib
    cache = load_embedding_cache()

    # Create hash of text for cache key
    text_hash = hashlib.md5(text.encode()).hexdigest()

    if text_hash in cache:
        return cache[text_hash]

    # Generate new embedding using load balancer
    if use_load_balancer:
        embedding_result = load_balancer.embed_distributed(EMBEDDING_MODEL, text, keep_alive=EMBEDDING_KEEP_ALIVE)
    else:
        embedding_result = ollama.embed(model=EMBEDDING_MODEL, input=text, keep_alive=EMBEDDING_KEEP_ALIVE)

    embeddings = embedding_result.embeddings if hasattr(embedding_result, 'embeddings') else []
    embedding = embeddings[0] if embeddings else []

    # Cache it
    cache[text_hash] = embedding
    save_embedding_cache(cache)

    return embedding

def load_document_index():
    """Load the document index or create it if it doesn't exist."""
    if not INDEX_FILE.exists():
        return {"documents": []}
    
    try:
        with open(INDEX_FILE, 'r') as f:
            return json.load(f)
    except (json.JSONDecodeError, FileNotFoundError):
        print(f"⚠️ Error loading index file. Creating a new one.")
        return {"documents": []}

def save_document_index(index_data):
    """Save the document index to disk."""
    with open(INDEX_FILE, 'w') as f:
        json.dump(index_data, f, indent=4)
    print(f"✅ Document index updated with {len(index_data['documents'])} documents")

def register_document(pdf_path, txt_path, content, chunks=None):
    """Register a processed document in the knowledge base index."""
    # Load existing index
    index_data = load_document_index()
    
    # Create document record
    document_id = f"doc_{len(index_data['documents']) + 1}"
    
    # Generate embeddings and chunks for search
    chunks = chunks or chunk_text(content)
    chunk_embeddings = []
    
    # Batch process embeddings for better performance
    print(f"🔄 Processing {len(chunks)} chunks in batches...")

    import hashlib
    cache = load_embedding_cache()
    uncached_chunks = []
    uncached_indices = []

    # Check cache first
    for i, chunk in enumerate(chunks):
        text_hash = hashlib.md5(chunk.encode()).hexdigest()
        if text_hash not in cache:
            uncached_chunks.append(chunk)
            uncached_indices.append(i)

    # Batch embed uncached chunks using load balancer
    if uncached_chunks:
        print(f"🚀 Embedding {len(uncached_chunks)} new chunks in parallel...")
        print(f"   Using {len(load_balancer.instances)} Ollama nodes")

        # Process in batches of 100 to save cache periodically
        batch_size = 100
        for batch_start in range(0, len(uncached_chunks), batch_size):
            batch_end = min(batch_start + batch_size, len(uncached_chunks))
            batch = uncached_chunks[batch_start:batch_end]

            print(f"   Processing batch {batch_start//batch_size + 1}/{(len(uncached_chunks) + batch_size - 1)//batch_size}...")

            batch_results = load_balancer.embed_batch(EMBEDDING_MODEL, batch)

            # Cache the batch embeddings
            cached_count = 0
            for chunk, result in zip(batch, batch_results):
                if result:
                    text_hash = hashlib.md5(chunk.encode()).hexdigest()
                    embeddings = result.embeddings if hasattr(result, 'embeddings') else []
                    embedding = embeddings[0] if embeddings else []
                    cache[text_hash] = embedding
                    cached_count += 1

            # Save cache after each batch
            save_embedding_cache(cache)
            print(f"   ✅ Cached {cached_count} embeddings from this batch")

        print(f"✅ All {len(uncached_chunks)} new embeddings cached")
    else:
        print(f"✅ All chunks found in cache!")

    # Now process all chunks
    for i, chunk in enumerate(chunks):
        try:
            # Show progress every 50 chunks
            if i % 50 == 0 and i > 0:
                print(f"🔄 Processed {i}/{len(chunks)} chunks...")

            # Get embedding from cache
            text_hash = hashlib.md5(chunk.encode()).hexdigest()
            embedding = cache.get(text_hash, [])

            # Store chunk with its embedding
            chunk_file = KB_DIR / f"{document_id}_chunk_{i}.json"
            chunk_data = {
                "text": chunk,
                "embedding": embedding
            }

            with open(chunk_file, 'w') as f:
                json.dump(chunk_data, f)

            # Remember the chunk reference
            chunk_embeddings.append({
                "chunk_id": f"{document_id}_chunk_{i}",
                "file": str(chunk_file)
            })
        except Exception as e:
            print(f"⚠️ Error embedding chunk {i}: {e}")
    
    # Add document to index
    doc_entry = {
        "id": document_id,
        "original": str(pdf_path),
        "text_path": str(txt_path),
        "processed_date": datetime.now().isoformat(),
        "chunks": chunk_embeddings
    }
    
    index_data["documents"].append(doc_entry)
    save_document_index(index_data)
    return document_id

def chunk_text(text, chunk_size=512, overlap=100):
    """
    Split text into overlapping chunks with intelligent token-aware splitting.

    Args:
        chunk_size: Target chunk size in tokens (approximate via chars * 0.25)
        overlap: Number of characters to overlap between chunks
    """
    # Token limits for mxbai-embed-large: 512 tokens max
    # Rough estimate: 1 token ≈ 4 characters
    MAX_TOKENS = 480  # Leave buffer for model
    MAX_CHARS = MAX_TOKENS * 4  # ~1920 chars
    TARGET_CHARS = chunk_size * 4  # ~2048 chars for chunk_size=512

    def split_large_text(text, max_size):
        """Recursively split text that's too large."""
        if len(text) <= max_size:
            return [text]

        # Try splitting by sentences first
        sentences = text.replace('! ', '!|').replace('? ', '?|').replace('. ', '.|').split('|')

        chunks = []
        current = []
        current_len = 0

        for sent in sentences:
            sent = sent.strip()
            if not sent:
                continue

            # If single sentence exceeds limit, split by words
            if len(sent) > max_size:
                words = sent.split()
                # Calculate words per chunk (with safety margin)
                words_per_chunk = int((max_size / len(sent)) * len(words) * 0.9)
                words_per_chunk = max(50, words_per_chunk)  # At least 50 words

                for i in range(0, len(words), words_per_chunk):
                    word_chunk = ' '.join(words[i:i + words_per_chunk])
                    if word_chunk:
                        chunks.append(word_chunk)
                continue

            # Add sentence to current chunk
            if current_len + len(sent) > max_size and current:
                chunks.append(' '.join(current))
                current = [sent]
                current_len = len(sent)
            else:
                current.append(sent)
                current_len += len(sent)

        if current:
            chunks.append(' '.join(current))

        return chunks

    # Split into paragraphs first
    paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]

    chunks = []
    current_chunk = []
    current_length = 0

    for para in paragraphs:
        para_len = len(para)

        # If paragraph is too large, split it first
        if para_len > MAX_CHARS:
            # Finalize current chunk if any
            if current_chunk:
                chunks.append('\n\n'.join(current_chunk))
                current_chunk = []
                current_length = 0

            # Split the large paragraph
            para_chunks = split_large_text(para, MAX_CHARS)
            chunks.extend(para_chunks)
            continue

        # Check if adding this paragraph exceeds target size
        if current_length + para_len > TARGET_CHARS and current_chunk:
            # Finalize current chunk
            chunks.append('\n\n'.join(current_chunk))

            # Start new chunk with overlap (keep last paragraph if small enough)
            if overlap > 0 and current_chunk and len(current_chunk[-1]) < overlap:
                current_chunk = [current_chunk[-1], para]
                current_length = len(current_chunk[-1]) + para_len
            else:
                current_chunk = [para]
                current_length = para_len
        else:
            current_chunk.append(para)
            current_length += para_len

    # Add final chunk
    if current_chunk:
        final_chunk = '\n\n'.join(current_chunk)
        # Safety check
        if len(final_chunk) > MAX_CHARS:
            chunks.extend(split_large_text(final_chunk, MAX_CHARS))
        else:
            chunks.append(final_chunk)

    # Final validation: ensure no chunk exceeds MAX_CHARS
    validated_chunks = []
    for chunk in chunks:
        if len(chunk) > MAX_CHARS:
            validated_chunks.extend(split_large_text(chunk, MAX_CHARS))
        else:
            validated_chunks.append(chunk)

    return validated_chunks

def list_documents():
    """List all processed documents in the knowledge base."""
    index_data = load_document_index()
    if not index_data["documents"]:
        print("📚 No documents have been processed yet.")
        return
    
    print(f"\n📚 Knowledge Base: {len(index_data['documents'])} documents")
    print("-" * 60)
    for i, doc in enumerate(index_data["documents"]):
        print(f"{i+1}. {Path(doc['original']).name}")
        print(f"   ID: {doc['id']} | Processed: {doc['processed_date'][:10]}")
        print(f"   Chunks: {len(doc['chunks'])}")
        print("-" * 60)

def get_similar_chunks(query, top_k=None, min_similarity=None):
    """Find text chunks similar to the query using vector similarity with adaptive top-k."""
    # Use configured defaults if not specified
    if min_similarity is None:
        min_similarity = RETRIEVAL_MIN_SIMILARITY

    try:
        # Get embedding for the query from cache
        query_embedding = get_cached_embedding(query)

        if not query_embedding:
            print("⚠️ Failed to generate query embedding")
            return []

        # Load document index
        index_data = load_document_index()

        # Check if we have documents
        if not index_data["documents"]:
            print("📚 No documents in knowledge base yet")
            return []

        # Adaptive top-k based on total chunks in database
        if top_k is None:
            total_chunks = sum(len(doc["chunks"]) for doc in index_data["documents"])

            # Scale top_k based on database size
            if total_chunks < 50:
                adaptive_k = min(total_chunks, 5)  # Very small DB, use fewer
            elif total_chunks < 200:
                adaptive_k = 10  # Small-medium DB, use default
            elif total_chunks < 1000:
                adaptive_k = 20  # Medium DB, retrieve more context
            else:
                adaptive_k = 30  # Large DB, need more chunks for good coverage

            top_k = adaptive_k
            print(f"   📊 Adaptive top-k: {top_k} (from {total_chunks} total chunks)")
        else:
            print(f"   📊 Using fixed top-k: {top_k}")

        # Collect all chunks with their embeddings
        chunks_with_similarity = []

        for doc in index_data["documents"]:
            for chunk_ref in doc["chunks"]:
                try:
                    # Load chunk data
                    chunk_file = Path(chunk_ref["file"])
                    if chunk_file.exists():
                        with open(chunk_file, 'r') as f:
                            chunk_data = json.load(f)

                        # Calculate cosine similarity
                        chunk_embedding = chunk_data.get("embedding", [])
                        if chunk_embedding:
                            similarity = cosine_similarity(query_embedding, chunk_embedding)

                            # Only include chunks above minimum similarity threshold
                            if similarity >= min_similarity:
                                chunks_with_similarity.append({
                                    "doc_id": doc["id"],
                                    "doc_name": Path(doc["original"]).name,
                                    "text": chunk_data["text"],
                                    "similarity": similarity
                                })
                except Exception as e:
                    print(f"⚠️ Error processing chunk {chunk_ref['chunk_id']}: {e}")

        # Sort by similarity (highest first) and get top k
        chunks_with_similarity.sort(key=lambda x: x["similarity"], reverse=True)

        # Return top k results
        results = chunks_with_similarity[:top_k]

        # Print retrieval stats
        print(f"   Found {len(results)} relevant chunks (similarity >= {min_similarity:.2f})")

        return results

    except Exception as e:
        print(f"⚠️ Error searching knowledge base: {e}")
        return []

def sanitize_for_xml(text):
    """Remove null bytes and control characters that break XML/DOCX."""
    import re
    # Remove NULL bytes
    text = text.replace('\x00', '')
    # Remove other control characters except newline, carriage return, and tab
    text = re.sub(r'[\x00-\x08\x0B-\x0C\x0E-\x1F\x7F-\x9F]', '', text)
    return text

def cosine_similarity(vec1, vec2):
    """Calculate cosine similarity between two vectors."""
    if not vec1 or not vec2:
        return 0

    vec1 = np.array(vec1)
    vec2 = np.array(vec2)

    dot_product = np.dot(vec1, vec2)
    norm_a = np.linalg.norm(vec1)
    norm_b = np.linalg.norm(vec2)

    if norm_a == 0 or norm_b == 0:
        return 0

    return dot_product / (norm_a * norm_b)

def embed_text(text):
    """Embeds text using Ollama without storing vector data in files."""
    try:
        # Using 'input' instead of 'prompt'
        response = ollama.embed(model=EMBEDDING_MODEL, input=text)
        return text  # Return the original text for saving to files
    except Exception as e:
        print(f"❌ Embedding error: {e}")
        return None

def extract_text_from_pdf(pdf_path):
    """Extracts text from a PDF file using multiple methods for better reliability."""
    pdf_path_str = str(pdf_path)
    extracted_text = ""
    
    # Method 1: Try PyPDF2 first
    try:
        print("🔍 Attempting extraction with PyPDF2...")
        reader = PdfReader(pdf_path_str)
        pypdf_text = ""
        
        for page_num, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text:
                pypdf_text += f"{page_text}\n\n"
            else:
                print(f"⚠️ PyPDF2: No text extracted from page {page_num + 1}")
        
        if pypdf_text.strip():
            print(f"✅ PyPDF2 successfully extracted {len(pypdf_text)} characters")
            extracted_text = pypdf_text
        else:
            print("⚠️ PyPDF2 extraction yielded no text, trying alternative method...")
    except Exception as e:
        print(f"⚠️ PyPDF2 extraction error: {e}")
    
    # Method 2: If PyPDF2 failed or returned no text, try pdftotext if available
    if not extracted_text:
        try:
            print("🔍 Attempting extraction with pdftotext (if installed)...")
            with tempfile.NamedTemporaryFile(suffix='.txt') as temp:
                # Try to use pdftotext (from poppler-utils) if installed
                result = subprocess.run(
                    ['pdftotext', '-layout', pdf_path_str, temp.name],
                    capture_output=True,
                    text=True
                )

                if result.returncode == 0:
                    with open(temp.name, 'r', encoding='utf-8') as f:
                        pdftotext_text = f.read()

                    if pdftotext_text.strip():
                        print(f"✅ pdftotext successfully extracted {len(pdftotext_text)} characters")
                        extracted_text = pdftotext_text
                    else:
                        print("⚠️ pdftotext extraction yielded no text")
                else:
                    print(f"⚠️ pdftotext error: {result.stderr}")
        except FileNotFoundError:
            print("⚠️ pdftotext not found on system, skipping alternative extraction")
        except Exception as e:
            print(f"⚠️ Alternative extraction error: {e}")

    # Method 3: If still no text, try OCR (for scanned documents)
    if not extracted_text or len(extracted_text.strip()) < 100:
        try:
            print("🔍 Attempting OCR extraction (for scanned/image-based PDFs)...")
            from pdf2image import convert_from_path
            import pytesseract

            # Convert PDF to images
            images = convert_from_path(pdf_path_str, dpi=300)
            print(f"📄 Converted PDF to {len(images)} image(s)")

            ocr_text = ""
            for i, image in enumerate(images, 1):
                print(f"   OCR processing page {i}/{len(images)}...")
                page_text = pytesseract.image_to_string(image, lang='eng')
                if page_text.strip():
                    ocr_text += f"--- Page {i} ---\n\n{page_text.strip()}\n\n"

            if ocr_text.strip():
                print(f"✅ OCR successfully extracted {len(ocr_text)} characters")
                extracted_text = ocr_text
            else:
                print("⚠️ OCR extraction yielded no text")

        except ImportError:
            print("⚠️ OCR libraries not available (pdf2image, pytesseract)")
            print("   Install with: pip install pdf2image pytesseract")
            print("   Also need: sudo apt-get install tesseract-ocr poppler-utils")
        except Exception as e:
            print(f"⚠️ OCR extraction error: {e}")

    # Check if we have any text after trying all methods
    if not extracted_text:
        print("❌ Failed to extract text with all available methods")
        return ""
    
    # Process the text to make it more readable
    processed_text = ""
    pages = extracted_text.split("\f")  # Form feed character often separates PDF pages
    
    for page_num, page_content in enumerate(pages):
        if page_content.strip():
            processed_text += f"--- Page {page_num + 1} ---\n\n{page_content.strip()}\n\n"
    
    return processed_text.strip()

def process_pdf(pdf_path):
    """Extracts text from PDF, embeds it, and saves clean conversions."""
    start_time = time.time()

    pdf_path = Path(pdf_path).resolve()
    if not pdf_path.exists():
        print(f"❌ Error: File not found → {pdf_path}")
        return

    print(f"📄 Processing '{pdf_path.name}'...")

    # Extract text from PDF using multiple methods
    extracted_text = extract_text_from_pdf(pdf_path)

    if not extracted_text:
        print(f"❌ Failed to extract text from {pdf_path.name}")
        print("💡 This PDF might be:")
        print("   - Scanned (image-based) without OCR")
        print("   - Protected/encrypted")
        print("   - Using non-standard fonts")
        print("   - Corrupted or malformed")
        return

    # Debug: Show a sample of the extracted text
    sample_length = min(200, len(extracted_text))
    print(f"📊 Extracted {len(extracted_text)} characters from {pdf_path.name}")
    print(f"📃 Sample of extracted text: \n{extracted_text[:sample_length]}...")

    # Prepare a clean version of text without the page markers for better readability
    clean_text = extracted_text
    if "--- Page" in clean_text:
        # Remove page markers if present but preserve content
        clean_text = "\n\n".join([
            line for line in clean_text.split("\n")
            if not line.strip().startswith("--- Page")
        ])

    # Get the PDF filename without extension for use in the document title and headings
    pdf_filename = pdf_path.stem
    
    # Save TXT - preserving the original filename
    txt_path = PROCESSED_DIR / f"{pdf_filename}.txt"
    with open(txt_path, "w", encoding="utf-8") as txt_file:
        # Add the PDF filename as the first line of the text file
        txt_file.write(f"# {pdf_filename}\n\n")
        txt_file.write(clean_text)
    print(f"✅ Saved TXT → {txt_path}")

    # Save Markdown - preserving the original filename
    md_path = PROCESSED_DIR / f"{pdf_filename}.md"
    with open(md_path, "w", encoding="utf-8") as md_file:
        # Creating proper markdown with the PDF filename as the title
        md_content = f"# {pdf_filename}\n\n{clean_text}"
        md_file.write(md_content)
    print(f"✅ Saved Markdown → {md_path}")

    # Save DOCX - preserving the original filename
    docx_path = PROCESSED_DIR / f"{pdf_filename}.docx"
    doc = docx.Document()

    # Add the PDF filename as the document title/heading
    doc.add_heading(sanitize_for_xml(pdf_filename), level=1)

    # Split text into paragraphs for better DOCX formatting
    paragraphs = clean_text.split('\n\n')
    for para in paragraphs:
        if para.strip():
            # Sanitize text to remove control characters that break XML
            sanitized_para = sanitize_for_xml(para.strip())
            if sanitized_para:
                doc.add_paragraph(sanitized_para)

    doc.save(docx_path)
    print(f"✅ Saved DOCX → {docx_path}")

    # Save JSON - preserving the original filename with metadata
    json_path = PROCESSED_DIR / f"{pdf_filename}.json"
    json_data = {
        "filename": pdf_filename,
        "original_path": str(pdf_path),
        "processed_date": datetime.now().isoformat(),
        "character_count": len(clean_text),
        "word_count": len(clean_text.split()),
        "title": pdf_filename,
        "content": clean_text,
        "metadata": {
            "extraction_method": "PyPDF2/pdftotext",
            "file_size_bytes": pdf_path.stat().st_size,
            "formats_generated": ["txt", "md", "docx", "json"]
        }
    }
    with open(json_path, "w", encoding="utf-8") as json_file:
        json.dump(json_data, json_file, indent=2, ensure_ascii=False)
    print(f"✅ Saved JSON → {json_path}")

    # Add to knowledge base for chat capability
    print("🧠 Adding document to knowledge base...")
    chunks = chunk_text(clean_text)
    print(f"📊 Document divided into {len(chunks)} semantic chunks")
    
    doc_id = register_document(pdf_path, txt_path, clean_text, chunks)
    print(f"✅ Document registered with ID: {doc_id}")

    elapsed_time = time.time() - start_time
    print(f"🎯 Completed processing {pdf_path.name}")
    print(f"⏱️  Total time: {elapsed_time:.2f}s ({elapsed_time/60:.1f} minutes)")

def process_directory(dir_path):
    """Processes all PDFs in a given directory."""
    start_time = time.time()

    dir_path = Path(dir_path).expanduser().resolve()

    print(f"🔍 Debug: Checking path → {dir_path}")

    if not dir_path.exists() or not dir_path.is_dir():
        print(f"❌ Error: Directory not found → {dir_path}")
        return

    pdf_files = list(dir_path.glob("*.pdf"))
    if not pdf_files:
        print(f"⚠️ No PDFs found in {dir_path}")
        return

    print(f"📂 Found {len(pdf_files)} PDFs. Processing...")

    for pdf in pdf_files:
        process_pdf(pdf)

    elapsed_time = time.time() - start_time
    print("✅ All PDFs processed!")
    print(f"⏱️  Total batch time: {elapsed_time:.2f}s ({elapsed_time/60:.1f} minutes)")
    print(f"📊 Average: {elapsed_time/len(pdf_files):.2f}s per PDF")

def chat():
    """Starts an interactive chat with embedded documents."""
    index_data = load_document_index()
    if not index_data["documents"]:
        print("📚 No documents in knowledge base yet. Process a PDF first.")
        return
    
    print("\n💬 Chat with your Documents")
    print("Type 'exit' to return to main menu")
    print(f"Knowledge base contains {len(index_data['documents'])} documents")
    
    chat_history = []
    
    while True:
        user_query = input("\n🙋 You: ").strip()
        
        if user_query.lower() == 'exit':
            print("Returning to main menu...")
            break
        
        if not user_query:
            continue
        
        # Start timing
        response_start_time = time.time()

        # Find relevant document chunks
        print("🔍 Searching knowledge base...")
        retrieval_start = time.time()
        relevant_chunks = get_similar_chunks(user_query)
        retrieval_time = time.time() - retrieval_start

        if not relevant_chunks:
            print("❌ No relevant information found in the knowledge base.")
            continue

        print(f"   ⏱️  Retrieval: {retrieval_time:.2f}s")

        # Document-aware intelligent context fitting
        # Conservative token limits for 2048-4096 context window models
        # Token estimation: 1 token ≈ 3.5 chars (conservative)
        # Reserve: system prompt (~100) + query (~150) + response (~1500) + history (~250)
        # Available for context: ~1500 tokens base, adjusted dynamically
        BASE_CONTEXT_TOKENS = 1500

        def estimate_tokens(text):
            """Conservative token estimation: 1 token ≈ 3.5 chars."""
            return int(len(text) / 3.5)

        # Group chunks by document
        from collections import defaultdict
        doc_chunks = defaultdict(list)
        for chunk in relevant_chunks:
            doc_chunks[chunk['doc_name']].append(chunk)

        num_docs = len(doc_chunks)
        print(f"   📚 Chunks span {num_docs} document(s)")

        # Dynamic strategy based on document count
        if num_docs == 1:
            # Single document: prioritize depth - use more chunks from same doc
            MAX_CONTEXT_TOKENS = BASE_CONTEXT_TOKENS * 1.3  # 1950 tokens
            min_chunks_per_doc = 3
            print(f"   🎯 Strategy: Deep dive (single document)")
        elif num_docs <= 3:
            # Few documents: balanced approach
            MAX_CONTEXT_TOKENS = BASE_CONTEXT_TOKENS * 1.1  # 1650 tokens
            min_chunks_per_doc = 2
            print(f"   🎯 Strategy: Balanced coverage ({num_docs} documents)")
        else:
            # Many documents: prioritize breadth - sample from each doc
            MAX_CONTEXT_TOKENS = BASE_CONTEXT_TOKENS  # 1500 tokens
            min_chunks_per_doc = 1
            print(f"   🎯 Strategy: Broad coverage ({num_docs} documents)")

        # Build context with document-aware selection
        context_parts = []
        current_tokens = 0
        chunks_used = 0
        docs_included = set()

        # Phase 1: Ensure minimum representation from each document
        for doc_name, chunks in sorted(doc_chunks.items(), key=lambda x: max(c['similarity'] for c in x[1]), reverse=True):
            doc_chunks_added = 0
            for chunk in sorted(chunks, key=lambda x: x['similarity'], reverse=True):
                if doc_chunks_added >= min_chunks_per_doc:
                    break

                chunk_text = chunk['text']
                similarity = chunk['similarity']
                formatted_chunk = f"[Doc: {doc_name}, Relevance: {similarity:.2f}]\n{chunk_text}"
                chunk_tokens = estimate_tokens(formatted_chunk)

                if current_tokens + chunk_tokens <= MAX_CONTEXT_TOKENS:
                    context_parts.append((similarity, formatted_chunk))
                    current_tokens += chunk_tokens
                    chunks_used += 1
                    doc_chunks_added += 1
                    docs_included.add(doc_name)

        # Phase 2: Fill remaining space with highest relevance chunks
        remaining_chunks = [
            chunk for doc_name, chunks in doc_chunks.items()
            for chunk in chunks
        ]
        remaining_chunks.sort(key=lambda x: x['similarity'], reverse=True)

        for chunk in remaining_chunks:
            if chunks_used >= len(relevant_chunks):
                break

            chunk_text = chunk['text']
            doc_name = chunk['doc_name']
            similarity = chunk['similarity']
            formatted_chunk = f"[Doc: {doc_name}, Relevance: {similarity:.2f}]\n{chunk_text}"

            # Skip if already included
            if any(formatted_chunk in part[1] for part in context_parts):
                continue

            chunk_tokens = estimate_tokens(formatted_chunk)

            if current_tokens + chunk_tokens <= MAX_CONTEXT_TOKENS:
                context_parts.append((similarity, formatted_chunk))
                current_tokens += chunk_tokens
                chunks_used += 1
                docs_included.add(doc_name)
            else:
                break

        # Sort final context by relevance
        context_parts.sort(key=lambda x: x[0], reverse=True)
        context = "\n\n".join([part[1] for part in context_parts])

        # Show fitting summary
        print(f"   📄 Selected {chunks_used}/{len(relevant_chunks)} chunks from {len(docs_included)} document(s) (~{current_tokens} tokens)")

        # Build prompt with context and chat history
        system_prompt = "You are FlockParser AI, a helpful assistant that answers questions based on the user's documents. Only use information from the provided document context. If you don't know or the answer isn't in the context, say so."

        # Build user message with context and optional history
        user_message_parts = []

        if chat_history:
            history_text = "\n".join([
                f"Previous Q: {q}\nPrevious A: {a}"
                for q, a in chat_history[-2:]  # Last 2 exchanges only
            ])
            user_message_parts.append(f"CHAT HISTORY:\n{history_text}\n")

        user_message_parts.append(f"CONTEXT FROM DOCUMENTS:\n{context}")
        user_message_parts.append(f"\nQUESTION: {user_query}")
        user_message = "\n".join(user_message_parts)

        # Generate response using LLM with load balancing
        print("🤖 Generating response...")
        generation_start = time.time()
        try:
            response = load_balancer.chat_distributed(
                model=CHAT_MODEL,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_message}
                ],
                keep_alive=CHAT_KEEP_ALIVE
            )

            generation_time = time.time() - generation_start
            answer = response['message']['content']

            # Display response
            print(f"\n🤖 AI: {answer}")

            # Update chat history
            chat_history.append((user_query, answer))

            # Show source documents
            print("\n📚 Sources:")
            for i, chunk in enumerate(relevant_chunks[:CHUNKS_TO_SHOW]):
                print(f"  {i+1}. {chunk['doc_name']} (relevance: {chunk['similarity']:.2f})")

            # Show timing breakdown
            total_time = time.time() - response_start_time
            print(f"\n⏱️  Response timing:")
            print(f"   Retrieval: {retrieval_time:.2f}s")
            print(f"   Generation: {generation_time:.2f}s")
            print(f"   Total: {total_time:.2f}s")

        except Exception as e:
            print(f"❌ Error generating response: {e}")

def check_dependencies():
    """Checks for the presence of external tools that might help with PDF processing."""
    print("🔍 Checking for helpful dependencies...")
    
    # Check for pdftotext (from Poppler utils)
    try:
        result = subprocess.run(['pdftotext', '-v'], 
                               stdout=subprocess.PIPE, 
                               stderr=subprocess.PIPE, 
                               text=True)
        if result.returncode == 0:
            version_info = result.stderr.strip() if result.stderr else result.stdout.strip()
            print(f"✅ pdftotext found: {version_info}")
        else:
            print("❌ pdftotext is not working properly")
    except FileNotFoundError:
        print("❌ pdftotext not found. For better PDF extraction, consider installing:")
        print("   - Linux: sudo apt-get install poppler-utils")
        print("   - macOS: brew install poppler")
        print("   - Windows: Install from http://blog.alivate.com.au/poppler-windows/")
    
    # Check PyPDF2 version
    import pkg_resources
    try:
        pypdf_version = pkg_resources.get_distribution("PyPDF2").version
        print(f"✅ PyPDF2 version: {pypdf_version}")
    except pkg_resources.DistributionNotFound:
        print("❌ PyPDF2 not found in installed packages")
    
    # Check for OCRmyPDF for potential enhancement
    try:
        result = subprocess.run(['ocrmypdf', '--version'], 
                               stdout=subprocess.PIPE, 
                               stderr=subprocess.PIPE, 
                               text=True)
        if result.returncode == 0:
            version_info = result.stdout.strip()
            print(f"✅ OCRmyPDF found: {version_info}")
            print("   This can be used to add OCR to scanned PDFs if needed")
        else:
            print("❌ OCRmyPDF is not working properly")
    except FileNotFoundError:
        print("ℹ️ OCRmyPDF not found (optional for OCR capability)")
    
    # Check Ollama availability
    try:
        result = ollama.list()
        # Display available models - handle both dict and object response formats
        if hasattr(result, 'models'):
            # New API returns object with models attribute
            models = result.models
            model_names = [model.model if hasattr(model, 'model') else str(model) for model in models]
        else:
            # Fallback for dict format
            models = result.get('models', [])
            model_names = [model.get('name', 'unknown') for model in models]

        # Check if embedding model exists (with or without :latest suffix)
        embedding_found = any(EMBEDDING_MODEL in name for name in model_names)
        if embedding_found:
            print(f"✅ Embedding model '{EMBEDDING_MODEL}' is available")
        else:
            print(f"⚠️ Embedding model '{EMBEDDING_MODEL}' not found in Ollama")
            print(f"   Run 'ollama pull {EMBEDDING_MODEL}' to download it")

        # Check if chat model exists (with or without :latest suffix)
        chat_found = any(CHAT_MODEL in name for name in model_names)
        if chat_found:
            print(f"✅ Chat model '{CHAT_MODEL}' is available")
        else:
            print(f"⚠️ Chat model '{CHAT_MODEL}' not found in Ollama")
            print(f"   Run 'ollama pull {CHAT_MODEL}' to download it")
            
    except Exception as e:
        print(f"❌ Ollama not available or error connecting: {e}")
        print("   Make sure Ollama is installed and running")
        
    print("\n💡 Missing tools can be installed to improve PDF processing capabilities")

def clear_cache():
    """Clear the embedding cache."""
    try:
        if EMBEDDING_CACHE_FILE.exists():
            confirm = input("⚠️  This will delete the embedding cache. Continue? (yes/no): ").strip().lower()
            if confirm == "yes":
                EMBEDDING_CACHE_FILE.unlink()
                print("✅ Embedding cache cleared successfully")
                print("   Next PDF processing will regenerate embeddings")
            else:
                print("❌ Operation cancelled")
        else:
            print("ℹ️ No embedding cache found")
    except Exception as e:
        print(f"❌ Error clearing cache: {e}")

def gpu_status():
    """Show intelligent GPU routing status."""
    print("\n" + "=" * 70)
    print("🎯 INTELLIGENT GPU ROUTING STATUS")
    print("=" * 70)

    router = IntelligentGPURouter(load_balancer.instances)
    router.print_cluster_report()

def gpu_route_model(model_name: str):
    """Show routing decision for a specific model."""
    print("\n" + "=" * 70)
    print(f"🧠 ROUTING DECISION FOR: {model_name}")
    print("=" * 70)

    router = IntelligentGPURouter(load_balancer.instances)
    decision = router.route_model(model_name)

    print(f"\n📍 Recommended routing:")
    print(f"   Node: {decision['node']}")
    print(f"   Target: {decision['target']}")
    print(f"   Reason: {decision['reason']}")

def gpu_optimize():
    """Trigger intelligent GPU optimization."""
    priority_models = [EMBEDDING_MODEL, CHAT_MODEL]

    print("\n" + "=" * 70)
    print(f"🔧 OPTIMIZING {len(priority_models)} PRIORITY MODELS")
    print("=" * 70)

    router = IntelligentGPURouter(load_balancer.instances)
    router.optimize_cluster(priority_models)

def gpu_check_fit(model_name: str):
    """Check which nodes can fit a specific model."""
    print("\n" + "=" * 70)
    print(f"✅ CHECKING FIT FOR: {model_name}")
    print("=" * 70)

    router = IntelligentGPURouter(load_balancer.instances)
    model_size = router.get_model_size(model_name)

    print(f"\n📦 Model size: {model_size} MB")
    print("\n📍 Node compatibility:")

    for node_url in load_balancer.instances:
        can_fit, reason = router.can_fit_on_gpu(node_url, model_name)
        if can_fit:
            print(f"   ✅ {node_url}: {reason}")
        else:
            print(f"   ❌ {node_url}: {reason}")

def gpu_list_models():
    """List all known models and their sizes."""
    print("\n" + "=" * 70)
    print("📚 KNOWN MODELS DATABASE")
    print("=" * 70)

    router = IntelligentGPURouter(load_balancer.instances)

    print("\n📦 Model sizes:")
    for model, size_mb in sorted(router.known_model_sizes.items(), key=lambda x: x[1]):
        size_gb = size_mb / 1024
        print(f"   {model:30s} {size_mb:6.0f} MB ({size_gb:.2f} GB)")

def unload_model(model_name: str):
    """Unload a specific model from all nodes."""
    print(f"\n🗑️  Unloading {model_name} from all nodes...")

    for node_url in load_balancer.instances:
        try:
            # Use keep_alive=0 to unload immediately
            if 'embed' in model_name.lower():
                response = requests.post(
                    f"{node_url}/api/embed",
                    json={
                        "model": model_name,
                        "input": "unload",
                        "keep_alive": 0
                    },
                    timeout=10
                )
            else:
                response = requests.post(
                    f"{node_url}/api/generate",
                    json={
                        "model": model_name,
                        "prompt": "unload",
                        "keep_alive": 0
                    },
                    timeout=10
                )

            if response.status_code == 200:
                print(f"   ✅ {node_url}: Unloaded {model_name}")
            else:
                print(f"   ⚠️  {node_url}: Status {response.status_code}")

        except Exception as e:
            print(f"   ⚠️  {node_url}: {str(e)}")

    print(f"\n✅ Unload requests sent for {model_name}")

def cleanup_models():
    """Unload all non-priority models from all nodes."""
    priority_models = {EMBEDDING_MODEL, CHAT_MODEL}

    print("\n🧹 Cleaning up non-priority models...")
    print(f"   Priority models: {', '.join(priority_models)}")

    models_to_unload = set()

    # Check what's loaded on each node
    for node_url in load_balancer.instances:
        try:
            response = requests.get(f"{node_url}/api/ps", timeout=5)
            if response.status_code == 200:
                data = response.json()
                for model in data.get('models', []):
                    model_name = model.get('name', '')
                    # Check if this model is NOT a priority model
                    is_priority = any(priority in model_name for priority in priority_models)
                    if not is_priority and model_name:
                        models_to_unload.add(model_name)
        except Exception as e:
            print(f"   ⚠️  Error checking {node_url}: {e}")

    if not models_to_unload:
        print("\n✅ No non-priority models to unload")
        return

    print(f"\n📋 Found {len(models_to_unload)} non-priority models:")
    for model in models_to_unload:
        print(f"   - {model}")

    confirm = input("\n⚠️  Unload these models? (yes/no): ").strip().lower()
    if confirm != "yes":
        print("❌ Operation cancelled")
        return

    # Unload each model
    for model_name in models_to_unload:
        unload_model(model_name)

    print("\n✅ Cleanup complete!")

def clear_db():
    """Clear the ChromaDB vector store (removes all documents)."""
    try:
        confirm = input("⚠️  This will DELETE ALL DOCUMENTS from the vector database. Continue? (yes/no): ").strip().lower()
        if confirm != "yes":
            print("❌ Operation cancelled")
            return

        global chroma_collection

        # Delete the collection
        try:
            chroma_client.delete_collection(name="documents")
            print("✅ ChromaDB collection deleted")
        except Exception:
            pass  # Collection might not exist

        # Recreate the collection
        chroma_collection = chroma_client.get_or_create_collection(
            name="documents",
            metadata={"hnsw:space": "cosine"}
        )
        print("✅ ChromaDB vector store cleared successfully")
        print("   All documents removed from database")

        # Optionally clear the document index too
        clear_index = input("Also clear document index? (yes/no): ").strip().lower()
        if clear_index == "yes":
            if INDEX_FILE.exists():
                INDEX_FILE.unlink()
            print("✅ Document index cleared")

        # Optionally clear JSON knowledge base
        clear_json = input("Also clear legacy JSON knowledge base? (yes/no): ").strip().lower()
        if clear_json == "yes":
            json_files = list(KB_DIR.glob("*.json"))
            for f in json_files:
                f.unlink()
            print(f"✅ Cleared {len(json_files)} JSON files from knowledge base")

    except Exception as e:
        print(f"❌ Error clearing database: {e}")

def vram_report():
    """Show detailed VRAM usage report for all nodes."""
    monitor = VRAMMonitor()

    print(f"\n🔍 Detected GPU type: {monitor.gpu_type.upper() if monitor.gpu_type != 'none' else 'None (CPU only)'}")

    # Get local report
    print("\n📊 Local Node Report:")
    local_report = monitor.get_comprehensive_report("http://localhost:11434")
    monitor.print_report(local_report)

    # Get distributed nodes report
    if len(load_balancer.instances) > 1:
        print("\n🌐 Distributed Nodes Report:")
        node_results = monitor_distributed_nodes(load_balancer.instances)

        for node_url, info in node_results.items():
            if info['status'] == 'online':
                gpu_status = "🚀 GPU" if info['gpu_accelerated'] else "🐢 CPU"
                vram_gb = info['vram_mb'] / 1024
                ram_gb = info['ram_mb'] / 1024

                print(f"\n   {gpu_status} {node_url}:")
                if info['gpu_accelerated']:
                    print(f"      VRAM Usage: {vram_gb:.2f} GB")
                else:
                    print(f"      RAM Usage: {ram_gb:.2f} GB (CPU fallback)")

                if info['models']:
                    print(f"      Loaded Models:")
                    for model in info['models']:
                        print(f"         - {model['name']} ({model['location']})")
            else:
                print(f"   ❌ {node_url}: {info['error']}")

        print("\n" + "="*70)

def main():
    """Command-line interface."""
    print("🚀 Welcome to FlockParser")
    print(COMMANDS)
    
    # Quick dependency check on startup
    print("\nℹ️ Run 'check_deps' for detailed dependency information")

    while True:
        command = input("\n⚡ Enter command: ").strip().split()

        if not command:
            continue

        action = command[0]
        arg = " ".join(command[1:]) if len(command) > 1 else None

        if action == "open_pdf" and arg:
            process_pdf(arg)
        elif action == "open_dir" and arg:
            process_directory(arg)
        elif action == "chat":
            chat()
        elif action == "list_docs":
            list_documents()
        elif action == "check_deps":
            check_dependencies()
        elif action == "discover_nodes":
            load_balancer.discover_nodes()
        elif action == "add_node" and arg:
            load_balancer.add_node(arg)
        elif action == "remove_node" and arg:
            load_balancer.remove_node(arg)
        elif action == "list_nodes":
            load_balancer.list_nodes()
        elif action == "verify_models":
            load_balancer.verify_models_on_nodes()
        elif action == "lb_stats":
            load_balancer.print_stats()
        elif action == "set_routing" and arg:
            load_balancer.set_routing_strategy(arg)
        elif action == "vram_report":
            vram_report()
        elif action == "force_gpu" and arg:
            load_balancer.force_gpu_all_nodes(arg)
        elif action == "gpu_status":
            gpu_status()
        elif action == "gpu_route" and arg:
            gpu_route_model(arg)
        elif action == "gpu_optimize":
            gpu_optimize()
        elif action == "gpu_check" and arg:
            gpu_check_fit(arg)
        elif action == "gpu_models":
            gpu_list_models()
        elif action == "unload_model" and arg:
            unload_model(arg)
        elif action == "cleanup_models":
            cleanup_models()
        elif action == "parallelism_report":
            from adaptive_parallelism import print_parallelism_report
            print_parallelism_report(load_balancer)
        elif action == "clear_cache":
            clear_cache()
        elif action == "clear_db":
            clear_db()
        elif action == "exit":
            print("👋 Exiting. See you next time!")
            load_balancer.print_stats()  # Show stats on exit
            break
        else:
            print("⚠️ Invalid command. Try again.")
            print(COMMANDS)

if __name__ == "__main__":
    main()
