from __future__ import annotations
from typing import Iterable
from ..types import UnifiedDocument, Metadata, Section, Chunk
from ..core import register_parser, ParserProtocol

class DOCXParser(ParserProtocol):
    name = "docx"
    content_types: Iterable[str] = ("application/vnd.openxmlformats-officedocument.wordprocessingml.document",)
    extensions: Iterable[str] = (".docx",)

    def can_parse(self, meta: Metadata) -> bool:
        return (meta.content_type in self.content_types or 
                (meta.path or "").endswith(".docx"))

    def parse(self, target, meta: Metadata, recursive: bool = False, **kwargs) -> UnifiedDocument:
        doc = UnifiedDocument(meta=meta, sections=[])
        
        try:
            from docx import Document
            
            # Load the document
            document = Document(str(target))
            
            # Extract document properties
            if document.core_properties.title:
                doc.meta.title = document.core_properties.title
            if document.core_properties.author:
                doc.meta.extra["author"] = document.core_properties.author
            if document.core_properties.subject:
                doc.meta.extra["subject"] = document.core_properties.subject
            if document.core_properties.created:
                doc.meta.extra["created"] = document.core_properties.created.isoformat()
            if document.core_properties.modified:
                doc.meta.extra["modified"] = document.core_properties.modified.isoformat()
            
            # Extract paragraphs
            paragraphs = []
            current_section = None
            current_chunks = []
            order = 0
            
            for paragraph in document.paragraphs:
                text = paragraph.text.strip()
                if not text:
                    continue
                
                # Check if this is a heading (simple heuristic)
                if (paragraph.style.name.startswith('Heading') or 
                    len(text) < 100 and text.isupper() or
                    paragraph.style.name in ['Title', 'Subtitle']):
                    
                    # Save previous section
                    if current_section and current_chunks:
                        doc.sections.append(Section(
                            heading=current_section,
                            chunks=current_chunks,
                            meta={"type": "content"}
                        ))
                    
                    # Start new section
                    current_section = text
                    current_chunks = []
                    order += 1
                else:
                    # Add to current section
                    chunk = Chunk(text=text, order=order, meta={"type": "paragraph"})
                    current_chunks.append(chunk)
                    order += 1
            
            # Save final section
            if current_section and current_chunks:
                doc.sections.append(Section(
                    heading=current_section,
                    chunks=current_chunks,
                    meta={"type": "content"}
                ))
            elif current_chunks:
                doc.sections.append(Section(
                    heading="Content",
                    chunks=current_chunks,
                    meta={"type": "content"}
                ))
            
            # Extract tables
            for i, table in enumerate(document.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(" | ".join(row_data))
                
                if table_data:
                    table_text = "\n".join(table_data)
                    doc.sections.append(Section(
                        heading=f"Table {i + 1}",
                        chunks=[Chunk(text=table_text, order=order, meta={"type": "table"})],
                        meta={"type": "table", "table_index": i}
                    ))
                    order += 1
            
            # Store document structure info
            doc.meta.extra["document_info"] = {
                "paragraph_count": len(document.paragraphs),
                "table_count": len(document.tables),
                "has_images": len(document.inline_shapes) > 0
            }
            
        except ImportError:
            error_text = "[panparsex:docx] python-docx not available. Install with: pip install python-docx"
            doc.sections.append(Section(
                heading="Error",
                chunks=[Chunk(text=error_text, order=0, meta={"error": True})],
                meta={"error": True}
            ))
        except Exception as e:
            error_text = f"[panparsex:docx] unable to parse: {e}"
            doc.sections.append(Section(
                heading="Error",
                chunks=[Chunk(text=error_text, order=0, meta={"error": True})],
                meta={"error": True}
            ))
        
        return doc

register_parser(DOCXParser())
