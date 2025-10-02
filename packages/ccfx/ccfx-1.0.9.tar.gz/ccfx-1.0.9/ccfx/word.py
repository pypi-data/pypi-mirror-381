#!/bin/python3

'''
a module to easily create microsoft word documents from python

Author  : Celray James CHAWANDA
Email   : celray.chawanda@outlook.com
Licence : MIT 2023
Repo    : https://github.com/celray

Date    : 2023-07-20
'''

# imports
import os
from docx import Document
from docx.shared import Inches, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# classes
class word_document:
    def __init__(self, path) -> None:
        self.document = Document()
        self.path     = path if path.endswith('docx') else f'{path}.docx'
    
    def addHeading(self, heading, level = 2):
        self.document.add_heading(heading, level)
    
    def addParagraph(self, text = "", alignment = 'justify'):
        '''
        allignments:
                - justify,
                - left,
                - center,
                - right,
                - justify-low,
                - justify-med
        '''
        
        p = self.document.add_paragraph(text)
        
        if alignment == 'left':p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        if alignment == 'center':p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        if alignment == 'right':p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        if alignment == 'justify':p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        if alignment == 'justify-low':p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY_LOW
        if alignment == 'justify-med':p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY_MED


    def addListItem(self, text = "", numbers = False):
        p = self.document.add_paragraph(text)
        if numbers:
            p.style = 'List Number'
        else:
            p.style = 'List Bullet'
    
    def addText(self, text, bold = False, italic = False):
        if bold:
            self.document.paragraphs[-1].add_run(text).bold = True
        elif italic:
            self.document.paragraphs[-1].add_run(text).italic = True
        elif bold and italic:
            self.document.paragraphs[-1].add_run(text).bold = True
            self.document.paragraphs[-1].runs[-1].italic = True
        else:
            self.document.paragraphs[-1].add_run(text)


    def addImage(self, path_to_image, width_ = 16):
        self.document.add_picture(path_to_image, width=Cm(width_))

    def addPageBreak(self):
        self.document.add_page_break()

    def save(self):
        self.create_path(self.path)
        self.document.save(self.path)

    def setMargins(self, margin = 1.75):
        sections = self.document.sections
        for section in sections:
            section.top_margin = Cm(margin)
            section.bottom_margin = Cm(margin)
            section.left_margin = Cm(margin)
            section.right_margin = Cm(margin)

    def createPath(self, path_name, v = False):
        path_name = os.path.dirname(path_name)
        if path_name == '':
            path_name = './'
        if not os.path.isdir(path_name):
            os.makedirs(path_name)
            if v:
                print(f"\t> created path: {path_name}")
        
        return path_name
